"""
Document Processing MCP Server
Handles secure document parsing and AI-powered content extraction
"""

import os
import json
import logging
from typing import Dict, List, Any, Optional
from pathlib import Path
from datetime import datetime

import PyPDF2
from docx import Document
from anthropic import Anthropic
import mcp
from mcp.server.fastmcp import FastMCP

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Anthropic client
claude_client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

# Initialize FastMCP server
mcp = FastMCP(
    "DocumentProcessor",
    instructions="Document processing server for resume text extraction and role identification",
    host="localhost", 
    port=8002,
)

@mcp.tool()
async def extract_text_from_file(file_path: str, file_type: str) -> Dict[str, Any]:
    """
    Extract raw text from supported file formats.
    
    Args:
        file_path (str): Path to the file to process
        file_type (str): File extension (pdf, docx, doc, txt)
    
    Returns:
        Dict containing extracted text and metadata
    """
    
    result = {
        "success": False,
        "text": "",
        "error": "",
        "metadata": {
            "file_path": file_path,
            "file_type": file_type,
            "extraction_timestamp": datetime.now().isoformat(),
            "page_count": 0,
            "character_count": 0
        }
    }
    
    try:
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            result["error"] = f"File not found: {file_path}"
            return result
        
        extracted_text = ""
        
        if file_type.lower() == 'pdf':
            extracted_text, page_count = _extract_from_pdf(file_path)
            result["metadata"]["page_count"] = page_count
            
        elif file_type.lower() in ['docx', 'doc']:
            extracted_text = _extract_from_docx(file_path)
            
        elif file_type.lower() == 'txt':
            extracted_text = _extract_from_txt(file_path)
        
        else:
            result["error"] = f"Unsupported file type: {file_type}"
            return result
        
        # Basic text validation
        if not extracted_text or len(extracted_text.strip()) < 10:
            result["error"] = "Insufficient text content extracted"
            return result
        
        result.update({
            "success": True,
            "text": extracted_text,
            "metadata": {
                **result["metadata"],
                "character_count": len(extracted_text),
                "word_count": len(extracted_text.split()),
                "lines_count": len(extracted_text.splitlines())
            }
        })
        
        logger.info(f"Successfully extracted {len(extracted_text)} characters from {file_path}")
        
    except Exception as e:
        result["error"] = f"Text extraction failed: {str(e)}"
        logger.error(f"Text extraction failed for {file_path}: {e}")
    
    return result

@mcp.tool()
async def extract_professional_roles(text: str, document_id: str, client_name: str = "unknown") -> Dict[str, Any]:
    """
    Use Claude Sonnet 4 to extract structured professional history from resume text.
    
    Args:
        text (str): Raw text from the resume
        document_id (str): Unique identifier for the document
        client_name (str): Name of the resume owner
    
    Returns:
        Dict containing extracted roles and metadata
    """
    
    result = {
        "success": False,
        "document_id": document_id,
        "client_name": client_name,
        "roles": [],
        "extraction_metadata": {},
        "error": ""
    }
    
    try:
        # Comprehensive extraction prompt for Claude
        extraction_prompt = """
        You are an expert at extracting professional role information from resumes. 
        Extract ALL professional roles from this resume text, including internships, part-time roles, and consulting positions.
        
        For each role, provide the following information in JSON format:
        
        REQUIRED FIELDS (always include, use null if not found):
        - company: Company/organization name
        - title: Job title/position
        - start_year: Start year (YYYY format, number)
        - end_year: End year (YYYY format, number) or null if current position
        - start_month: Start month (1-12, number) or null if not specified
        - end_month: End month (1-12, number) or null if not specified
        
        DETAILED INFORMATION (extract if available):
        - manager_title: Title of direct manager/supervisor
        - direct_reports: Array of direct report titles/functions
        - budget_responsibility: Budget managed (number in USD) or null
        - headcount: Team size managed (number) or null
        - quota: Sales/revenue quota (number in USD) or null
        - peer_functions: Array of peer departments/functions worked with
        - achievements: Array of key achievements with metrics
        - responsibilities: Array of main job responsibilities
        - location: Work location (city, state/country)
        - employment_type: "full-time", "part-time", "contract", "internship", etc.
        
        CONFIDENCE AND SOURCING:
        - confidence_score: Your confidence in the extraction (0.0 to 1.0)
        - source_indicators: Array of text snippets that support each field
        
        Return ONLY a valid JSON array of role objects. No additional text or formatting.
        Be very careful with dates - extract exactly what's written.
        If employment is listed as "2020 - Present" or similar, set end_year to null.
        
        Resume text to analyze:
        {text}
        """
        
        # Call Claude for extraction
        response = claude_client.messages.create(
            model="claude-3-5-sonnet-20241022",
            temperature=0.1,  # Low temperature for consistent extraction
            max_tokens=8000,
            messages=[{
                "role": "user",
                "content": extraction_prompt.format(text=text)
            }]
        )
        
        # Parse Claude's response
        response_text = response.content[0].text.strip()
        
        # Try to extract JSON from the response
        try:
            # Remove any markdown formatting if present
            if response_text.startswith('```json'):
                response_text = response_text.replace('```json', '').replace('```', '').strip()
            elif response_text.startswith('```'):
                response_text = response_text.replace('```', '').strip()
            
            extracted_roles = json.loads(response_text)
            
            # Validate the extracted data
            if not isinstance(extracted_roles, list):
                raise ValueError("Response is not a list of roles")
            
            # Post-process and validate each role
            validated_roles = []
            for i, role in enumerate(extracted_roles):
                if not isinstance(role, dict):
                    continue
                
                # Ensure required fields exist
                validated_role = {
                    "company": role.get("company", "").strip() if role.get("company") else "",
                    "title": role.get("title", "").strip() if role.get("title") else "",
                    "start_year": role.get("start_year"),
                    "end_year": role.get("end_year"),
                    "start_month": role.get("start_month"),
                    "end_month": role.get("end_month"),
                    "manager_title": role.get("manager_title"),
                    "direct_reports": role.get("direct_reports", []),
                    "budget_responsibility": role.get("budget_responsibility"),
                    "headcount": role.get("headcount"),
                    "quota": role.get("quota"),
                    "peer_functions": role.get("peer_functions", []),
                    "achievements": role.get("achievements", []),
                    "responsibilities": role.get("responsibilities", []),
                    "location": role.get("location"),
                    "employment_type": role.get("employment_type"),
                    "confidence_score": role.get("confidence_score", 0.7),
                    "source_indicators": role.get("source_indicators", []),
                    "role_index": i
                }
                
                # Skip roles missing critical information
                if not validated_role["company"] or not validated_role["title"]:
                    logger.warning(f"Skipping role {i} due to missing company or title")
                    continue
                
                validated_roles.append(validated_role)
            
            result.update({
                "success": True,
                "roles": validated_roles,
                "extraction_metadata": {
                    "model": "claude-3-5-sonnet-20241022",
                    "extraction_timestamp": datetime.now().isoformat(),
                    "total_roles_found": len(validated_roles),
                    "average_confidence": sum(r.get("confidence_score", 0) for r in validated_roles) / len(validated_roles) if validated_roles else 0,
                    "text_length": len(text),
                    "prompt_tokens": response.usage.input_tokens,
                    "completion_tokens": response.usage.output_tokens
                }
            })
            
            logger.info(f"Successfully extracted {len(validated_roles)} roles from document {document_id}")
            
        except json.JSONDecodeError as e:
            result["error"] = f"Failed to parse Claude response as JSON: {str(e)}"
            logger.error(f"JSON parse error for document {document_id}: {e}")
            logger.debug(f"Claude response was: {response_text[:500]}...")
            
    except Exception as e:
        result["error"] = f"Role extraction failed: {str(e)}"
        logger.error(f"Role extraction failed for document {document_id}: {e}")
    
    return result

@mcp.tool()
async def enhance_role_data(role_data: Dict[str, Any], context_text: str) -> Dict[str, Any]:
    """
    Use Claude to enhance/enrich a specific role with additional context.
    
    Args:
        role_data (Dict): Basic role information to enhance
        context_text (str): Additional context from the resume
    
    Returns:
        Dict containing enhanced role data
    """
    
    try:
        enhancement_prompt = """
        Given this basic role information and additional context, enhance the role data with more details.
        Focus on extracting specific metrics, achievements, and responsibilities.
        
        Current role data:
        {role_json}
        
        Additional context:
        {context}
        
        Return the enhanced role as JSON with the same structure but more detailed information.
        Pay special attention to:
        - Quantified achievements (numbers, percentages, dollar amounts)
        - Team size and reporting structure
        - Budget or revenue responsibility
        - Key technologies or methodologies used
        - Impact metrics
        
        Return ONLY valid JSON, no additional text.
        """
        
        response = claude_client.messages.create(
            model="claude-3-5-sonnet-20241022",
            temperature=0.1,
            max_tokens=4000,
            messages=[{
                "role": "user",
                "content": enhancement_prompt.format(
                    role_json=json.dumps(role_data, indent=2),
                    context=context_text
                )
            }]
        )
        
        enhanced_data = json.loads(response.content[0].text.strip())
        return {"success": True, "enhanced_role": enhanced_data}
        
    except Exception as e:
        logger.error(f"Role enhancement failed: {e}")
        return {"success": False, "error": str(e), "original_role": role_data}

def _extract_from_pdf(file_path: str) -> tuple[str, int]:
    """Extract text from PDF file using PyPDF2."""
    text = ""
    page_count = 0
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
    
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise
    
    return text.strip(), page_count

def _extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX file using python-docx."""
    text = ""
    
    try:
        doc = Document(file_path)
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
    
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise
    
    return text.strip()

def _extract_from_txt(file_path: str) -> str:
    """Extract text from TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        # Try different encodings if UTF-8 fails
        if not text.strip():
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        if text.strip():
                            break
                except:
                    continue
    
    except Exception as e:
        logger.error(f"TXT extraction failed: {e}")
        raise
    
    return text.strip()

if __name__ == "__main__":
    logger.info("Starting Document Processor MCP Server...")
    mcp.run(transport="stdio") 