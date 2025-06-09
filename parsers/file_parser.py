import os
from typing import Union

try:
    from pypdf2 import PdfReader
except ImportError:
    # Fallback or error message if pypdf2 is not installed
    # This helps in environments where only some parsers are needed/installed
    PdfReader = None
    print("Warning: pypdf2 library not found. PDF parsing will not be available.")

try:
    from docx import Document
except ImportError:
    Document = None
    print("Warning: python-docx library not found. DOCX parsing will not be available.")

class FileParserError(Exception):
    """Custom exception for file parsing errors."""
    pass

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extracts text from all pages of a PDF file.

    Args:
        file_path: The path to the PDF file.

    Returns:
        The concatenated text extracted from the PDF.
        Returns an empty string if the PdfReader is not available or if an error occurs.

    Raises:
        FileParserError: If the file is not found or if there's an issue parsing the PDF.
    """
    if PdfReader is None:
        raise FileParserError("PdfReader is not available (pypdf2 library missing or failed to import). Cannot parse PDF.")

    if not os.path.exists(file_path):
        raise FileParserError(f"File not found: {file_path}")

    text_parts = []
    try:
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            if reader.is_encrypted:
                # Attempt to decrypt with an empty password, common for some PDFs.
                # For password-protected PDFs, a password would be needed here.
                try:
                    reader.decrypt('')
                except Exception as decrypt_err:
                    raise FileParserError(f"Could not decrypt PDF '{file_path}': {decrypt_err}")

            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text_parts.append(page.extract_text() or "") # Ensure None is handled
        return "\n".join(text_parts)
    except FileNotFoundError:
        raise FileParserError(f"File not found: {file_path}")
    except Exception as e:
        # Catching pypdf2 specific errors or other generic errors
        raise FileParserError(f"Error parsing PDF file '{file_path}': {e}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts text from all paragraphs of a DOCX file.

    Args:
        file_path: The path to the DOCX file.

    Returns:
        The concatenated text extracted from the DOCX.
        Returns an empty string if Document is not available or if an error occurs.

    Raises:
        FileParserError: If the file is not found or if there's an issue parsing the DOCX.
    """
    if Document is None:
        raise FileParserError("Document is not available (python-docx library missing or failed to import). Cannot parse DOCX.")

    if not os.path.exists(file_path):
        raise FileParserError(f"File not found: {file_path}")

    text_parts = []
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text_parts.append(para.text)
        return "\n".join(text_parts)
    except FileNotFoundError:
        raise FileParserError(f"File not found: {file_path}")
    except Exception as e:
        # Catches errors from python-docx or other issues
        raise FileParserError(f"Error parsing DOCX file '{file_path}': {e}")


def extract_text_from_txt(file_path: str) -> str:
    """
    Extracts text from a TXT file.

    Args:
        file_path: The path to the TXT file.

    Returns:
        The content of the TXT file.

    Raises:
        FileParserError: If the file is not found or if there's an issue reading the file.
    """
    if not os.path.exists(file_path):
        raise FileParserError(f"File not found: {file_path}")

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        raise FileParserError(f"File not found: {file_path}")
    except UnicodeDecodeError:
        # Try with a different encoding if UTF-8 fails, or raise specific error
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            raise FileParserError(f"Error reading TXT file '{file_path}' with UTF-8 and Latin-1 encodings: {e}")
    except Exception as e:
        raise FileParserError(f"Error reading TXT file '{file_path}': {e}")


if __name__ == "__main__":
    # Create a temporary directory for test data
    test_data_dir = "temp_test_data_parsers"
    os.makedirs(test_data_dir, exist_ok=True)

    # Define file paths
    sample_txt_path = os.path.join(test_data_dir, "sample.txt")
    sample_docx_path = os.path.join(test_data_dir, "sample.docx")
    sample_pdf_path = os.path.join(test_data_dir, "sample.pdf")

    # Create dummy files
    print(f"Creating dummy files in {os.path.abspath(test_data_dir)}...")
    with open(sample_txt_path, "w", encoding="utf-8") as f:
        f.write("This is a sample text file.\nIt has multiple lines.\nHello, World!")
    print(f"Created {sample_txt_path}")

    if Document is not None:
        try:
            doc = Document()
            doc.add_paragraph("This is a sample DOCX document.")
            doc.add_paragraph("It also has multiple paragraphs.")
            doc.add_paragraph("Created by python-docx for testing.")
            doc.save(sample_docx_path)
            print(f"Created {sample_docx_path}")
        except Exception as e:
            print(f"Could not create sample.docx: {e}. DOCX tests will be skipped.")
            sample_docx_path = None # Skip test if creation fails
    else:
        print("python-docx not available. Skipping sample.docx creation and test.")
        sample_docx_path = None

    # PDF creation is more complex, so we'll just test the error handling for a non-existent file
    # or rely on manual creation if a simple PDF library was available for writing.
    # For this test, we'll primarily focus on the TXT and DOCX if available.
    # A real PDF test would require a sample PDF file.
    print(f"For PDF testing, please manually create a '{sample_pdf_path}' or ensure pypdf2 can handle non-existent file errors gracefully.")


    print("\n--- Testing TXT Extraction ---")
    try:
        txt_content = extract_text_from_txt(sample_txt_path)
        print(f"Successfully extracted from TXT:\n---\n{txt_content}\n---")
    except FileParserError as e:
        print(f"Error extracting from TXT: {e}")
    except Exception as e:
        print(f"An unexpected error occurred during TXT extraction: {e}")

    if sample_docx_path and Document is not None:
        print("\n--- Testing DOCX Extraction ---")
        try:
            docx_content = extract_text_from_docx(sample_docx_path)
            print(f"Successfully extracted from DOCX:\n---\n{docx_content}\n---")
        except FileParserError as e:
            print(f"Error extracting from DOCX: {e}")
        except Exception as e:
            print(f"An unexpected error occurred during DOCX extraction: {e}")
    else:
        print("\n--- Skipping DOCX Extraction Test (library or file not available) ---")

    print("\n--- Testing PDF Extraction (Error Handling for Non-Existent File) ---")
    non_existent_pdf = os.path.join(test_data_dir, "non_existent.pdf")
    try:
        if PdfReader is not None:
            pdf_content = extract_text_from_pdf(non_existent_pdf)
            print(f"Extracted from PDF (this should not happen for a non-existent file):\n{pdf_content}")
        else:
            print("pypdf2 not available. Skipping PDF non-existent file test logic that uses PdfReader.")
            # Still, try to call the function to see if it raises the correct FileParserError due to missing PdfReader
            try:
                 extract_text_from_pdf(non_existent_pdf)
            except FileParserError as e:
                 print(f"Correctly caught expected error for PDF parsing without PdfReader: {e}")


    except FileParserError as e:
        print(f"Successfully caught expected error for non-existent PDF: {e}")
    except Exception as e:
        print(f"An unexpected error occurred during non-existent PDF test: {e}")

    # Example of testing with a manually created PDF (if you add one to test_data_dir)
    # if os.path.exists(sample_pdf_path) and PdfReader is not None:
    #     print("\n--- Testing PDF Extraction (With Sample File if Present) ---")
    #     try:
    #         pdf_content = extract_text_from_pdf(sample_pdf_path)
    #         print(f"Successfully extracted from PDF '{sample_pdf_path}':\n---\n{pdf_content}\n---")
    #     except FileParserError as e:
    #         print(f"Error extracting from PDF '{sample_pdf_path}': {e}")
    #     except Exception as e:
    #         print(f"An unexpected error occurred during PDF extraction from '{sample_pdf_path}': {e}")
    # else:
    #     print(f"\n--- Skipping PDF Extraction Test (Sample file '{sample_pdf_path}' not found or pypdf2 not available) ---")

    print(f"\nNote: For comprehensive PDF testing, place a valid 'sample.pdf' in the '{test_data_dir}' directory.")
    print("Test script finished. Consider removing the 'temp_test_data_parsers' directory if not needed.")
